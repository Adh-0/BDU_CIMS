"""
loader.py — Universal document loader.
Extracts clean text from PDF, DOCX, TXT, MD, and CSV files.
Focuses on extracting ACCURATE text — no garbled data, no missing content.
"""

import csv
import io
import logging
from pathlib import Path
from dataclasses import dataclass, field

import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """A single extracted document chunk with metadata."""
    content: str
    metadata: dict = field(default_factory=dict)


# ── PDF Extraction ───────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Clean extracted text — remove excessive whitespace, fix encoding."""
    if not text:
        return ""
    # Collapse multiple spaces/tabs into single space
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if line:
            # Collapse internal whitespace
            line = " ".join(line.split())
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def _table_to_markdown(table: list[list]) -> str:
    """Convert a pdfplumber table (list of rows) to clean markdown.
    Handles None cells, vertical text, and empty rows gracefully.
    """
    if not table or len(table) < 2:
        return ""

    # Clean all cells
    cleaned = []
    for row in table:
        cleaned_row = []
        for cell in row:
            text = str(cell) if cell is not None else ""
            text = text.strip()
            # Fix vertical text (single chars separated by \n)
            if "\n" in text and all(len(part.strip()) <= 2 for part in text.split("\n")):
                text = "".join(part.strip() for part in text.split("\n"))
            else:
                text = " ".join(text.split())
            cleaned_row.append(text)
        cleaned.append(cleaned_row)

    # Skip empty rows
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if len(cleaned) < 2:
        return ""

    # Build markdown table
    header = cleaned[0]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in cleaned[1:]:
        # Pad or trim row to match header length
        while len(row) < len(header):
            row.append("")
        row = row[: len(header)]
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _is_valid_text(text: str, min_length: int = 20) -> bool:
    """Check if extracted text is actually readable (not garbled).
    Rejects chunks that are mostly special characters or too short.
    """
    if not text or len(text.strip()) < min_length:
        return False

    # Count readable characters (letters, digits, common punctuation)
    readable = sum(1 for c in text if c.isalnum() or c in " .,;:!?-/()\n")
    ratio = readable / len(text)

    # If less than 60% readable characters, likely garbled
    return ratio >= 0.60


def load_pdf(filepath: Path) -> list[LoadedDocument]:
    """Extract text from PDF using pdfplumber.
    - Text pages: extract_text() for clean paragraph text.
    - Table pages: extract tables as markdown format.
    - Validates all extracted text for quality.
    """
    docs = []
    source = filepath.name

    try:
        with pdfplumber.open(filepath) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"  Processing PDF: {source} ({total_pages} pages)")

            for page_idx, page in enumerate(pdf.pages):
                page_num = page_idx + 1
                tables = page.extract_tables()

                # ── Extract table content ────────────────────────
                if tables:
                    for table_idx, table in enumerate(tables):
                        md_table = _table_to_markdown(table)
                        if md_table and _is_valid_text(md_table, min_length=30):
                            docs.append(LoadedDocument(
                                content=md_table,
                                metadata={
                                    "source": source,
                                    "page": page_num,
                                    "total_pages": total_pages,
                                    "content_type": "table",
                                    "file_type": "pdf",
                                },
                            ))

                # ── Extract text content ─────────────────────────
                # Always try to get text even from table pages
                # (there may be text outside the table regions)
                text = page.extract_text() or ""
                text = _clean_text(text)

                if _is_valid_text(text):
                    docs.append(LoadedDocument(
                        content=text,
                        metadata={
                            "source": source,
                            "page": page_num,
                            "total_pages": total_pages,
                            "content_type": "text",
                            "file_type": "pdf",
                        },
                    ))

    except Exception as e:
        logger.error(f"  ERROR reading PDF {source}: {e}")

    return docs


# ── DOCX Extraction ──────────────────────────────────────────────

def load_docx(filepath: Path) -> list[LoadedDocument]:
    """Extract text from Word documents (.docx).
    Handles paragraphs and tables separately.
    """
    docs = []
    source = filepath.name

    try:
        doc = DocxDocument(str(filepath))
        logger.info(f"  Processing DOCX: {source}")

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        if _is_valid_text(full_text):
            docs.append(LoadedDocument(
                content=full_text,
                metadata={
                    "source": source,
                    "content_type": "text",
                    "file_type": "docx",
                },
            ))

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            md_table = _table_to_markdown(rows)
            if md_table and _is_valid_text(md_table, min_length=30):
                docs.append(LoadedDocument(
                    content=md_table,
                    metadata={
                        "source": source,
                        "content_type": "table",
                        "file_type": "docx",
                    },
                ))

    except Exception as e:
        logger.error(f"  ERROR reading DOCX {source}: {e}")

    return docs


# ── Plain Text Extraction ────────────────────────────────────────

def load_text(filepath: Path) -> list[LoadedDocument]:
    """Load plain text files (.txt, .md)."""
    source = filepath.name

    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            text = filepath.read_text(encoding=encoding)
            text = _clean_text(text)
            if _is_valid_text(text):
                return [LoadedDocument(
                    content=text,
                    metadata={
                        "source": source,
                        "content_type": "text",
                        "file_type": filepath.suffix.lstrip("."),
                    },
                )]
            else:
                logger.warning(f"  SKIP {source}: text quality too low")
                return []
        except UnicodeDecodeError:
            continue

    logger.error(f"  ERROR reading {source}: could not decode with any encoding")
    return []


# ── CSV Extraction ───────────────────────────────────────────────

def load_csv(filepath: Path) -> list[LoadedDocument]:
    """Load CSV files as structured text.
    Each row becomes a readable text block using column headers.
    """
    docs = []
    source = filepath.name

    try:
        text = filepath.read_text(encoding="utf-8-sig")
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return []

        logger.info(f"  Processing CSV: {source} ({len(rows)} rows)")

        # Convert entire CSV to a markdown table
        headers = list(rows[0].keys())
        md_lines = ["| " + " | ".join(headers) + " |"]
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            values = [str(row.get(h, "")).strip() for h in headers]
            md_lines.append("| " + " | ".join(values) + " |")

        md_table = "\n".join(md_lines)
        if _is_valid_text(md_table, min_length=30):
            docs.append(LoadedDocument(
                content=md_table,
                metadata={
                    "source": source,
                    "content_type": "table",
                    "file_type": "csv",
                },
            ))

    except Exception as e:
        logger.error(f"  ERROR reading CSV {source}: {e}")

    return docs


# ── Main Loader ──────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
    ".csv": load_csv,
}


def load_document(filepath: Path) -> list[LoadedDocument]:
    """Load a single document based on its file extension."""
    ext = filepath.suffix.lower()
    loader_fn = SUPPORTED_EXTENSIONS.get(ext)

    if not loader_fn:
        logger.warning(f"  SKIP {filepath.name}: unsupported format ({ext})")
        return []

    return loader_fn(filepath)


def load_all_documents(directory: Path) -> list[LoadedDocument]:
    """Recursively load all supported documents from a directory.

    Returns:
        List of LoadedDocument objects with clean text and metadata.
    """
    all_docs = []

    if not directory.exists():
        logger.error(f"Documents directory not found: {directory}")
        return all_docs

    files = sorted(
        p for p in directory.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    if not files:
        logger.warning(f"No supported documents found in {directory}")
        return all_docs

    logger.info(f"Found {len(files)} supported files in {directory}")

    for filepath in files:
        try:
            docs = load_document(filepath)
            all_docs.extend(docs)
        except Exception as e:
            logger.error(f"  FATAL ERROR loading {filepath.name}: {e}")

    logger.info(f"Total extracted: {len(all_docs)} document segments")
    return all_docs
